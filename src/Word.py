import docx
from docx.enum.section import WD_ORIENT

# Create an instance of a word document
from Planning import *
from tkinter.filedialog import askdirectory
from docx.shared import RGBColor


def word(jour, nom, planning, theme_t, user, liste_eleves):
    if user == "Lena":
        lettre = "L"
    elif user == "Karine":
        lettre = "K"
    planning_t = [
        planning.planning,
        planning.ancien_planning,
        planning.ancien_planning2,
        planning.ancien_planning3,
    ]
    path = askdirectory()
    print(nom)
    nom = nom[::-1]
    planning_t = planning_t[::-1]
    theme_t = theme_t[::-1]
    for heure in planning.liste_eleve:
        liste_cavalier = [eleve for [eleve, nb] in liste_eleves[heure]]
        if lettre.lower() in heure.lower() or "semaine" in nom[0]:
            nomfichier = (
                nom[len(nom) - 1]
                .lower()
                .replace("liste samedi", "")
                .replace(".xlsx", "")
                .replace("liste mercredi", "")
                + "_"
                + heure
                + ".docx"
            )
            liste_eleve = set()
            for i, plan in enumerate(planning_t):
                for cellule in plan:
                    if cellule[0] == heure:
                        # print("la")
                        liste_eleve.add(cellule[2])
            # print(liste_eleve)
            doc = docx.Document()
            section = doc.sections[-1]
            new_width, new_height = section.page_height, section.page_width
            section.orientation = WD_ORIENT.LANDSCAPE
            section.page_width = new_width
            section.page_height = new_height

            table = doc.add_table(rows=len(liste_eleve) + 3, cols=5)
            table.style = "Table Grid"
            table.cell(0, 0).text = jour + "/" + heure

            # Adding heading in the 1st row of the table
            for i, date in enumerate(nom):
                date = (
                    date.lower()
                    .replace("liste samedi", "")
                    .replace(".xlsx", "")
                    .replace("liste mercredi", "")
                )
                table.cell(0, i + 1).text = date
            eleves = {}

            ind_rattrapage = []
            # Adding names of students in the first column
            for i, eleve in enumerate(liste_eleve):
                eleves[eleve] = i + 1
                table.cell(i + 1, 0).text = eleve
                if eleve not in liste_cavalier:
                    ind_rattrapage.append(i + 1)
                    paragraphe = table.cell(i + 1, 0).paragraphs[0]
                    paragraphe.text = eleve
                    paragraphe.runs[0].font.color.rgb = RGBColor(255, 0, 0)

            # Adding data to the table
            # print(planning_t)

            for i, plan in enumerate(planning_t):
                ind = 1
                for cellule in plan:
                    if cellule[0] == heure and cellule[2] in eleves:
                        # print(cellule)
                        table.cell(eleves[cellule[2]], i + 1).text = cellule[1]
                        # print(eleves[cellule[2]],ind_rattrapage)
                        if eleves[cellule[2]] in ind_rattrapage:
                            paragraphe = table.cell(
                                eleves[cellule[2]], i + 1
                            ).paragraphs[0]
                            paragraphe.text = cellule[1]
                            paragraphe.runs[0].font.color.rgb = RGBColor(255, 0, 0)
                        ind += 1
            table.cell(len(liste_eleve) + 1, 0).text = "theme"

            for i, theme in enumerate(theme_t):
                if heure in theme:
                    table.cell(len(liste_eleve) + 1, i + 1).text = theme[heure]
            for i in range(5):
                table.cell(len(liste_eleve) + 2, i).text = "\r\r\r\r\r\r\r\r\r\r"
            print(path + "/" + nomfichier)
            doc.save(path + "/" + nomfichier)
